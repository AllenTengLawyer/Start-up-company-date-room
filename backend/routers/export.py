from fastapi import APIRouter, HTTPException, Request, Query
from fastapi.responses import StreamingResponse, HTMLResponse, JSONResponse
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel
import os
import io
import re
import shutil
import json
import zipfile
from urllib.parse import quote

router = APIRouter()

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")

def get_report_data(project_id: int):
    from ..database import get_db
    from datetime import datetime
    try:
        from zoneinfo import ZoneInfo
        generated_at = datetime.now(ZoneInfo("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M")
    except Exception:
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    items = db.execute(
        "SELECT * FROM ldd_items WHERE project_id=? ORDER BY sort_order", (project_id,)
    ).fetchall()
    statuses = {
        r["ldd_item_id"]: dict(r)
        for r in db.execute("""
            SELECT ls.* FROM ldd_status ls
            JOIN ldd_items li ON ls.ldd_item_id = li.id
            WHERE li.project_id=?
        """, (project_id,)).fetchall()
    }
    mappings_raw = db.execute("""
        SELECT lm.ldd_item_id, f.file_name, COALESCE(lm.notes, '') AS notes
        FROM ldd_mappings lm
        JOIN files f ON lm.file_id = f.id
        JOIN ldd_items li ON lm.ldd_item_id = li.id
        WHERE li.project_id=?
    """, (project_id,)).fetchall()
    mappings_by_item = {}
    for m in mappings_raw:
        mappings_by_item.setdefault(m["ldd_item_id"], []).append({
            "file_name": m["file_name"],
            "notes": m["notes"] or ""
        })

    score_data = db.execute("""
        SELECT ls.status, COUNT(*) as n
        FROM ldd_status ls JOIN ldd_items li ON ls.ldd_item_id=li.id
        WHERE li.project_id=? AND li.is_required=1 GROUP BY ls.status
    """, (project_id,)).fetchall()
    sc = {r["status"]: r["n"] for r in score_data}
    total = db.execute("SELECT COUNT(*) as n FROM ldd_items WHERE project_id=? AND is_required=1", (project_id,)).fetchone()["n"]
    db.close()

    provided = sc.get("provided", 0)
    partial = sc.get("partial", 0)
    na = sc.get("na", 0)
    effective = total - na
    score_pct = round((provided + partial * 0.5) / effective * 100) if effective > 0 else 0

    sections = {}
    for item in items:
        d = dict(item)
        s = statuses.get(item["id"], {})
        d["status"] = s.get("status", "pending")
        d["statement"] = s.get("statement", "")
        d["files"] = mappings_by_item.get(item["id"], [])
        sec = item["section_no"]
        sections.setdefault(sec, {"section_no": sec, "items": []})
        sections[sec]["items"].append(d)

    return {
        "project": dict(project),
        "sections": list(sections.values()),
        "score_pct": score_pct,
        "total": total,
        "provided": provided,
        "partial": partial,
        "na": na,
        "pending": total - provided - partial - na,
        "generated_at": generated_at,
    }

@router.get("/projects/{project_id}/export/html")
def export_html(project_id: int, lang: str = Query("zh")):
    data = get_report_data(project_id)
    lang = "en" if str(lang).lower().startswith("en") else "zh"
    labels = {
        "zh": {
            "html_lang": "zh",
            "title": "DD就绪报告",
            "h1": "法律尽职调查就绪报告",
            "meta_company": "公司",
            "meta_generated_at": "生成时间",
            "meta_internal": "仅供内部使用",
            "score_label": "DD 就绪度评分",
            "stat_provided": "已提供",
            "stat_partial": "部分提供",
            "stat_pending": "未提供",
            "stat_na": "不适用",
            "th_no": "编号",
            "th_item": "事项",
            "th_risk": "风险",
            "th_status": "状态",
            "th_files": "已提供文件",
            "th_notes": "说明",
            "risk_high": "高",
            "risk_medium": "中",
            "risk_low": "低",
            "status_provided": "✓ 已提供",
            "status_partial": "△ 部分提供",
            "status_pending": "✗ 未提供",
            "status_na": "— 不适用",
            "item_statement": "条目说明",
        },
        "en": {
            "html_lang": "en",
            "title": "DD Readiness Report",
            "h1": "Due Diligence Readiness Report",
            "meta_company": "Company",
            "meta_generated_at": "Generated at",
            "meta_internal": "Internal use only",
            "score_label": "DD Readiness Score",
            "stat_provided": "Provided",
            "stat_partial": "Partial",
            "stat_pending": "Pending",
            "stat_na": "N/A",
            "th_no": "No",
            "th_item": "Item",
            "th_risk": "Risk",
            "th_status": "Status",
            "th_files": "Files",
            "th_notes": "Notes",
            "risk_high": "High",
            "risk_medium": "Medium",
            "risk_low": "Low",
            "status_provided": "✓ Provided",
            "status_partial": "△ Partial",
            "status_pending": "✗ Pending",
            "status_na": "— N/A",
            "item_statement": "Item note",
        },
    }[lang]
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))
    tmpl = env.get_template("report_template.html")
    html = tmpl.render(**data, lang=lang, labels=labels)
    return HTMLResponse(content=html)

@router.get("/projects/{project_id}/export/docx")
def export_docx(project_id: int, lang: str = Query("zh")):
    data = get_report_data(project_id)
    lang = "en" if str(lang).lower().startswith("en") else "zh"
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = Document()

        def set_east_asia(oxml_element, font_name: str):
            rPr = oxml_element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                oxml_element.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), font_name)

        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        set_east_asia(normal._element, "Microsoft YaHei")

        labels = {
            "zh": {
                "title": "法律尽职调查就绪报告",
                "generated_at": "生成时间",
                "section": "§",
                "th_no": "编号",
                "th_item": "事项",
                "th_risk": "风险",
                "th_status": "状态",
                "th_files": "已提供文件",
                "th_notes": "说明",
                "risk_high": "高",
                "risk_medium": "中",
                "risk_low": "低",
                "status_provided": "✓ 已提供",
                "status_partial": "△ 部分提供",
                "status_pending": "✗ 未提供",
                "status_na": "— 不适用",
                "item_statement": "条目说明",
                "metric": "指标",
                "value": "数值",
                "score": "就绪度评分",
                "required_items": "必需事项",
                "provided": "已提供",
                "partial": "部分提供",
                "pending": "未提供",
                "na": "不适用",
            },
            "en": {
                "title": "Due Diligence Readiness Report",
                "generated_at": "Generated at",
                "section": "Section ",
                "th_no": "No",
                "th_item": "Item",
                "th_risk": "Risk",
                "th_status": "Status",
                "th_files": "Files",
                "th_notes": "Notes",
                "risk_high": "High",
                "risk_medium": "Medium",
                "risk_low": "Low",
                "status_provided": "✓ Provided",
                "status_partial": "△ Partial",
                "status_pending": "✗ Pending",
                "status_na": "— N/A",
                "item_statement": "Item note",
                "metric": "Metric",
                "value": "Value",
                "score": "Readiness Score",
                "required_items": "Required Items",
                "provided": "Provided",
                "partial": "Partial",
                "pending": "Pending",
                "na": "N/A",
            },
        }[lang]

        title = doc.add_paragraph()
        r1 = title.add_run(labels["title"])
        r1.bold = True
        r1.font.size = Pt(22)
        r1.font.name = "Segoe UI"
        set_east_asia(r1._element, "Microsoft YaHei")

        project_name = (data.get("project") or {}).get("name") or ""
        if project_name:
            p2 = doc.add_paragraph()
            r2 = p2.add_run(str(project_name))
            r2.bold = True
            r2.font.size = Pt(14)
            set_east_asia(r2._element, "Microsoft YaHei")

        generated_at = data.get("generated_at") or ""
        if generated_at:
            p3 = doc.add_paragraph()
            r3 = p3.add_run(f"{labels['generated_at']}: {generated_at}")
            r3.font.size = Pt(10)
            set_east_asia(r3._element, "Microsoft YaHei")

        doc.add_paragraph()

        summary = doc.add_table(rows=1, cols=2)
        summary.style = "Table Grid"
        hdr = summary.rows[0].cells
        hdr[0].text = labels["metric"]
        hdr[1].text = labels["value"]
        rows = [
            (labels["score"], f"{data.get('score_pct', 0)}%"),
            (labels["required_items"], str(data.get("total", 0))),
            (labels["provided"], str(data.get("provided", 0))),
            (labels["partial"], str(data.get("partial", 0))),
            (labels["pending"], str(data.get("pending", 0))),
            (labels["na"], str(data.get("na", 0))),
        ]
        for k, v in rows:
            row = summary.add_row().cells
            row[0].text = k
            row[1].text = v

        doc.add_paragraph()

        status_map = {
            "provided": labels["status_provided"],
            "partial": labels["status_partial"],
            "pending": labels["status_pending"],
            "na": labels["status_na"],
        }

        risk_map = {
            "high": labels["risk_high"],
            "medium": labels["risk_medium"],
            "low": labels["risk_low"],
        }

        for sec in data.get("sections") or []:
            sec_no = sec.get("section_no")
            h = doc.add_paragraph()
            rh = h.add_run(f"{labels['section']}{sec_no}")
            rh.bold = True
            rh.font.size = Pt(14)
            set_east_asia(rh._element, "Microsoft YaHei")

            items = sec.get("items") or []
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style = "Table Grid"
            head = tbl.rows[0].cells
            head[0].text = labels["th_no"]
            head[1].text = labels["th_item"]
            head[2].text = labels["th_risk"]
            head[3].text = labels["th_status"]
            head[4].text = labels["th_files"]
            head[5].text = labels["th_notes"]
            for it in items:
                rr = tbl.add_row().cells
                rr[0].text = str(it.get("item_no") or "")
                rr[1].text = str(it.get("title") or it.get("name") or "")
                rr[2].text = risk_map.get(it.get("risk_level"), labels["risk_low"])
                rr[3].text = status_map.get(it.get("status"), str(it.get("status") or ""))

                files = it.get("files") or []
                rr[4].text = "\n".join([str(f.get("file_name") or "") for f in files if f])
                notes_lines = [str((f.get("notes") or "")).strip() for f in files if f]
                notes = "\n".join([n for n in notes_lines if n])
                statement = str((it.get("statement") or "")).strip()
                if statement:
                    notes = (notes + ("\n\n" if notes else "") + f"{labels['item_statement']}: {statement}")
                rr[5].text = notes

            doc.add_paragraph()

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return StreamingResponse(
            out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=DD_Report_{project_id}.docx"}
        )
    except Exception as e:
        raise HTTPException(500, f"Word生成失败：{str(e)}")


# ── Cabinet → Folder export ──────────────────────────────────────────────────

_ILLEGAL_CHARS = re.compile(r'[\\/:*?"<>|]')

def _sanitize(name: str) -> str:
    name = _ILLEGAL_CHARS.sub("_", name)
    return name.strip(". ") or "_"

def _content_disposition_attachment(filename: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("._") or "download"
    quoted = quote(filename, safe="")
    return f'attachment; filename="{safe}"; filename*=UTF-8\'\'{quoted}'

def _unique_dest(folder: str, filename: str) -> str:
    """Return a non-colliding destination path, appending _(2), _(3) etc."""
    dest = os.path.join(folder, filename)
    if not os.path.exists(dest):
        return dest
    base, ext = os.path.splitext(filename)
    n = 2
    while True:
        dest = os.path.join(folder, f"{base}_({n}){ext}")
        if not os.path.exists(dest):
            return dest
        n += 1

class FolderExportRequest(BaseModel):
    dest_path: str

@router.post("/projects/{project_id}/export/folder")
def export_folder(project_id: int, req: FolderExportRequest):
    from ..database import get_db
    dest = req.dest_path.strip()
    if not dest:
        raise HTTPException(400, "目标路径不能为空")

    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    root_path = project["root_path"]
    if os.path.normcase(os.path.abspath(dest)) == os.path.normcase(os.path.abspath(root_path)):
        db.close()
        raise HTTPException(400, "目标路径不能与源路径相同")

    # Build category_id → absolute folder path mapping
    cats = db.execute(
        "SELECT id, parent_id, name, sort_order FROM categories WHERE project_id=? ORDER BY sort_order",
        (project_id,)
    ).fetchall()
    id_to_cat = {r["id"]: dict(r) for r in cats}

    # Group children by parent to get sort-order index for prefix
    children_of = {}
    for c in cats:
        pid = c["parent_id"]
        children_of.setdefault(pid, []).append(c["id"])

    def folder_path(cat_id):
        parts = []
        cur = cat_id
        while cur is not None:
            cat = id_to_cat.get(cur)
            if cat is None:
                break
            siblings = children_of.get(cat["parent_id"], [])
            idx = siblings.index(cur) + 1 if cur in siblings else 1
            parts.append(f"{idx:02d}_{_sanitize(cat['name'])}")
            cur = cat["parent_id"]
        return os.path.join(dest, *reversed(parts))

    cat_paths = {cid: folder_path(cid) for cid in id_to_cat}

    # Load all registered files
    files = db.execute(
        "SELECT id, file_name, file_path, category_id FROM files WHERE project_id=?",
        (project_id,)
    ).fetchall()
    db.close()

    try:
        os.makedirs(dest, exist_ok=True)
    except PermissionError as e:
        raise HTTPException(400, f"无法创建目标目录：权限不足 ({dest})")

    copied = 0
    skipped = []

    for f in files:
        src = os.path.join(root_path, f["file_path"])
        if not os.path.isfile(src):
            skipped.append({"file_name": f["file_name"], "reason": "源文件不存在"})
            continue

        if f["category_id"] and f["category_id"] in cat_paths:
            target_folder = cat_paths[f["category_id"]]
        else:
            target_folder = os.path.join(dest, "_uncategorized")

        # Check path length (Windows MAX_PATH guard)
        candidate = os.path.join(target_folder, f["file_name"])
        if len(candidate) > 240:
            skipped.append({"file_name": f["file_name"], "reason": "目标路径超过系统长度限制"})
            continue

        try:
            os.makedirs(target_folder, exist_ok=True)
            dest_file = _unique_dest(target_folder, f["file_name"])
            shutil.copy2(src, dest_file)
            copied += 1
        except Exception as e:
            skipped.append({"file_name": f["file_name"], "reason": str(e)})

    return {"ok": True, "dest_path": dest, "copied": copied, "skipped": skipped}


# ── JSON backup / restore ─────────────────────────────────────────────────────

def _rows(db, sql, params=()):
    return [dict(r) for r in db.execute(sql, params).fetchall()]

@router.get("/projects/{project_id}/export/json")
def export_json(project_id: int):
    from ..database import get_db
    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    data = {
        "version": 1,
        "project": dict(project),
        "categories": _rows(db, "SELECT * FROM categories WHERE project_id=?", (project_id,)),
        "files": _rows(db, "SELECT * FROM files WHERE project_id=?", (project_id,)),
        "ldd_items": _rows(db, "SELECT * FROM ldd_items WHERE project_id=?", (project_id,)),
        "ldd_status": _rows(db, """
            SELECT ls.* FROM ldd_status ls
            JOIN ldd_items li ON ls.ldd_item_id = li.id
            WHERE li.project_id=?""", (project_id,)),
        "ldd_mappings": _rows(db, """
            SELECT lm.* FROM ldd_mappings lm
            JOIN ldd_items li ON lm.ldd_item_id = li.id
            WHERE li.project_id=?""", (project_id,)),
        "founders": _rows(db, "SELECT * FROM founders WHERE project_id=?", (project_id,)),
        "founder_checklist_status": _rows(db, """
            SELECT fcs.* FROM founder_checklist_status fcs
            JOIN founders f ON fcs.founder_id = f.id
            WHERE f.project_id=?""", (project_id,)),
        "founder_files": _rows(db, """
            SELECT ff.* FROM founder_files ff
            JOIN founders f ON ff.founder_id = f.id
            WHERE f.project_id=?""", (project_id,)),
    }
    db.close()

    project_name = dict(project)["name"].replace(" ", "_")
    safe_project = re.sub(r"[^A-Za-z0-9_.-]+", "_", project_name).strip("._") or f"project_{project_id}"
    filename = f"backup_{safe_project}_{project_id}.json"
    content = json.dumps(data, ensure_ascii=False, indent=2)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.post("/projects/import", status_code=201)
async def import_json(request: Request):
    from ..database import get_db
    try:
        data = await request.json()
    except Exception:
        raise HTTPException(400, "无效的JSON文件")

    if data.get("version") != 1:
        raise HTTPException(400, "不支持的备份版式")

    db = get_db()
    proj = data["project"]

    # Insert new project (new ID to avoid conflicts)
    cur = db.execute(
        "INSERT INTO projects (name, root_path, company_type, mode) VALUES (?,?,?,?)",
        (proj["name"] + " (导入)", proj["root_path"], proj.get("company_type", "cn"), proj.get("mode", "established"))
    )
    new_pid = cur.lastrowid

    # Map old IDs → new IDs for each table
    cat_map = {}
    for c in data.get("categories", []):
        r = db.execute(
            "INSERT INTO categories (project_id, parent_id, name, sort_order) VALUES (?,?,?,?)",
            (new_pid, None, c["name"], c.get("sort_order", 0))
        )
        cat_map[c["id"]] = r.lastrowid
    # Fix parent_id references
    for c in data.get("categories", []):
        if c.get("parent_id") and c["parent_id"] in cat_map:
            db.execute("UPDATE categories SET parent_id=? WHERE id=?",
                       (cat_map[c["parent_id"]], cat_map[c["id"]]))

    file_map = {}
    for f in data.get("files", []):
        r = db.execute(
            "INSERT INTO files (project_id, category_id, file_name, file_path, notes, keyword_suggested) VALUES (?,?,?,?,?,?)",
            (new_pid, cat_map.get(f.get("category_id")), f["file_name"], f["file_path"],
             f.get("notes", ""), f.get("keyword_suggested", 0))
        )
        file_map[f["id"]] = r.lastrowid

    ldd_map = {}
    for item in data.get("ldd_items", []):
        r = db.execute(
            """INSERT INTO ldd_items (project_id, section_no, item_no, title, title_en, description,
               item_type, risk_level, is_required, sort_order) VALUES (?,?,?,?,?,?,?,?,?,?)""",
            (new_pid, item["section_no"], item["item_no"], item["title"], item.get("title_en", ""),
             item.get("description", ""), item.get("item_type", "file"), item.get("risk_level", "medium"),
             item.get("is_required", 1), item.get("sort_order", 0))
        )
        ldd_map[item["id"]] = r.lastrowid

    for s in data.get("ldd_status", []):
        new_item_id = ldd_map.get(s["ldd_item_id"])
        if new_item_id:
            db.execute(
                "INSERT OR IGNORE INTO ldd_status (ldd_item_id, status, statement) VALUES (?,?,?)",
                (new_item_id, s["status"], s.get("statement", ""))
            )

    for m in data.get("ldd_mappings", []):
        new_item_id = ldd_map.get(m["ldd_item_id"])
        new_file_id = file_map.get(m["file_id"])
        if new_item_id and new_file_id:
            db.execute("INSERT OR IGNORE INTO ldd_mappings (ldd_item_id, file_id) VALUES (?,?)",
                       (new_item_id, new_file_id))

    founder_map = {}
    for f in data.get("founders", []):
        r = db.execute(
            "INSERT INTO founders (project_id, name, role, id_number, join_date, employment_type, notes) VALUES (?,?,?,?,?,?,?)",
            (new_pid, f["name"], f.get("role", ""), f.get("id_number", ""),
             f.get("join_date", ""), f.get("employment_type", "full_time"), f.get("notes", ""))
        )
        founder_map[f["id"]] = r.lastrowid

    for s in data.get("founder_checklist_status", []):
        new_fid = founder_map.get(s["founder_id"])
        if new_fid:
            db.execute(
                "INSERT OR IGNORE INTO founder_checklist_status (founder_id, item_code, status, statement) VALUES (?,?,?,?)",
                (new_fid, s["item_code"], s["status"], s.get("statement", ""))
            )

    for ff in data.get("founder_files", []):
        new_fid = founder_map.get(ff["founder_id"])
        if new_fid:
            db.execute(
                "INSERT INTO founder_files (founder_id, item_code, file_name, file_path, notes) VALUES (?,?,?,?,?)",
                (new_fid, ff["item_code"], ff["file_name"], ff["file_path"], ff.get("notes", ""))
            )

    db.commit()
    db.close()
    return {"id": new_pid, "name": proj["name"] + " (导入)"}


# ── LDD zip export ────────────────────────────────────────────────────────────

SECTION_NAMES = {
    '1': '集团公司基本文件', '2': '业务与重大合同', '3': '借款和担保',
    '4': '财务和会计', '5': '动产和不动产', '6': '知识产权',
    '7': '税务及财政补贴', '8': '雇员和不竞争', '9': '保险',
    '10': '诉讼、执行及行政处罚', '11': '网络安全、数据合规', '12': 'ESG', '13': '其他'
}

@router.get("/projects/{project_id}/export/ldd-zip")
def export_ldd_zip(project_id: int):
    from ..database import get_db
    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    root_path = project["root_path"]

    # Build category id → full path parts (for sub-folder hierarchy inside each LDD item)
    cats = db.execute(
        "SELECT id, parent_id, name FROM categories WHERE project_id=? ORDER BY sort_order",
        (project_id,)
    ).fetchall()
    id_to_cat = {r["id"]: dict(r) for r in cats}

    def cat_path_parts(cat_id):
        """Return list of category names from root to leaf."""
        parts = []
        cur = cat_id
        while cur is not None:
            cat = id_to_cat.get(cur)
            if cat is None:
                break
            parts.append(_sanitize(cat["name"]))
            cur = cat["parent_id"]
        return list(reversed(parts))

    # Load LDD items with their mappings
    items = db.execute(
        "SELECT * FROM ldd_items WHERE project_id=? ORDER BY sort_order", (project_id,)
    ).fetchall()

    mappings = db.execute("""
        SELECT lm.ldd_item_id, lm.notes, f.file_name, f.file_path, f.category_id,
               c.name as category_name, c.parent_id as category_parent_id
        FROM ldd_mappings lm
        JOIN files f ON lm.file_id = f.id
        JOIN ldd_items li ON lm.ldd_item_id = li.id
        LEFT JOIN categories c ON f.category_id = c.id
        WHERE li.project_id=?
    """, (project_id,)).fetchall()

    mappings_by_item = {}
    for m in mappings:
        mappings_by_item.setdefault(m["ldd_item_id"], []).append(dict(m))

    db.close()

    # Build zip in memory
    buf = io.BytesIO()
    seen_paths = {}  # track duplicates within zip

    def unique_zip_path(path):
        if path not in seen_paths:
            seen_paths[path] = 1
            return path
        seen_paths[path] += 1
        base, ext = os.path.splitext(path)
        return f"{base}_({seen_paths[path]}){ext}"

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Also add notes summary (statements + mapping notes)
        from ..database import get_db as _get_db2
        db2 = _get_db2()
        statuses = {
            r["ldd_item_id"]: dict(r)
            for r in db2.execute("""
                SELECT ls.* FROM ldd_status ls
                JOIN ldd_items li ON ls.ldd_item_id = li.id
                WHERE li.project_id=?
            """, (project_id,)).fetchall()
        }
        db2.close()
        lines = []
        lines.append(f"项目：{project['name']}")
        lines.append("")
        for item in items:
            sec = str(item["section_no"])
            sec_name = SECTION_NAMES.get(sec, sec)
            lines.append(f"§{sec} {sec_name} / {item['item_no']} {item['title']}")
            st = statuses.get(item["id"], {})
            if st.get("statement"):
                lines.append(f"说明：{st.get('statement')}")
            mfs = mappings_by_item.get(item["id"], [])
            for m in mfs:
                note = m.get("notes") or ""
                lines.append(f" - {m['file_name']}" + (f"（说明：{note}）" if note else ""))
            lines.append("")
        zf.writestr("LDD_Notes.txt", "\n".join(lines))

        for item in items:
            item_files = mappings_by_item.get(item["id"], [])
            if not item_files:
                continue

            sec = str(item["section_no"])
            sec_name = SECTION_NAMES.get(sec, sec)
            sec_folder = _sanitize(f"§{sec}_{sec_name}")
            item_folder = _sanitize(f"{item['item_no']}_{item['title']}")

            for f in item_files:
                src = os.path.join(root_path, f["file_path"])
                if not os.path.isfile(src):
                    continue

                # Build sub-path from category hierarchy
                if f["category_id"]:
                    sub_parts = cat_path_parts(f["category_id"])
                else:
                    sub_parts = []

                zip_parts = [sec_folder, item_folder] + sub_parts + [f["file_name"]]
                zip_path = unique_zip_path("/".join(zip_parts))

                try:
                    zf.write(src, zip_path)
                except Exception:
                    pass  # skip unreadable files

    buf.seek(0)
    project_name = _sanitize(project["name"])
    filename = f"LDD_{project_name}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": _content_disposition_attachment(filename)}
    )


# ── Category zip export ───────────────────────────────────────────────────────

@router.get("/projects/{project_id}/export/category-zip")
def export_category_zip(project_id: int, cat_id: int):
    from ..database import get_db
    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    root_path = project["root_path"]

    # Load all categories for this project
    cats = db.execute(
        "SELECT id, parent_id, name FROM categories WHERE project_id=? ORDER BY sort_order",
        (project_id,)
    ).fetchall()
    id_to_cat = {r["id"]: dict(r) for r in cats}

    # Collect all descendant category IDs (including cat_id itself)
    def descendant_ids(root_id):
        result = {root_id}
        changed = True
        while changed:
            changed = False
            for c in cats:
                if c["parent_id"] in result and c["id"] not in result:
                    result.add(c["id"])
                    changed = True
        return result

    target_ids = descendant_ids(cat_id)

    # Build relative folder path for a category (from cat_id root downward)
    def rel_folder(cid):
        parts = []
        cur = cid
        while cur is not None and cur != id_to_cat.get(cat_id, {}).get("parent_id"):
            cat = id_to_cat.get(cur)
            if cat is None:
                break
            parts.append(_sanitize(cat["name"]))
            cur = cat["parent_id"]
            if cur == id_to_cat[cat_id]["parent_id"]:
                break
        return "/".join(reversed(parts))

    # Load files in target categories
    placeholders = ",".join("?" * len(target_ids))
    files = db.execute(
        f"SELECT id, file_name, file_path, category_id FROM files WHERE project_id=? AND category_id IN ({placeholders})",
        (project_id, *target_ids)
    ).fetchall()
    db.close()

    buf = io.BytesIO()
    seen_paths = {}

    def unique_zip_path(path):
        if path not in seen_paths:
            seen_paths[path] = 1
            return path
        seen_paths[path] += 1
        base, ext = os.path.splitext(path)
        return f"{base}_({seen_paths[path]}){ext}"

    root_cat_name = _sanitize(id_to_cat[cat_id]["name"]) if cat_id in id_to_cat else str(cat_id)

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            src = os.path.join(root_path, f["file_path"])
            if not os.path.isfile(src):
                continue
            folder = rel_folder(f["category_id"]) if f["category_id"] else root_cat_name
            zip_path = unique_zip_path(f"{folder}/{f['file_name']}")
            try:
                zf.write(src, zip_path)
            except Exception:
                pass

    buf.seek(0)
    filename = f"{root_cat_name}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": _content_disposition_attachment(filename)}
    )


# ── Founder Background Report (Word + ZIP) ───────────────────────────────────

@router.get("/projects/{project_id}/export/founder-report")
def export_founder_report(project_id: int):
    from ..database import get_db
    from .founders import FOUNDER_CHECKLIST, DIMENSION_LABELS

    db = get_db()
    project = db.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        db.close()
        raise HTTPException(404, "项目不存在")

    founders = db.execute("SELECT * FROM founders WHERE project_id=? ORDER BY id", (project_id,)).fetchall()

    # collect checklist statuses and files per founder
    founder_data = []
    for founder in founders:
        fid = founder["id"]
        statuses = {
            r["item_code"]: dict(r)
            for r in db.execute("SELECT * FROM founder_checklist_status WHERE founder_id=?", (fid,)).fetchall()
        }
        files_by_code = {}
        all_files = db.execute("SELECT * FROM founder_files WHERE founder_id=?", (fid,)).fetchall()
        for f in all_files:
            code = f["item_code"] or ""
            files_by_code.setdefault(code, []).append(dict(f))

        items = []
        for item in FOUNDER_CHECKLIST:
            s = statuses.get(item["code"], {})
            items.append({
                **item,
                "status": s.get("status", "pending"),
                "statement": s.get("statement", ""),
                "files": files_by_code.get(item["code"], []),
            })

        founder_data.append({
            "founder": dict(founder),
            "items": items,
            "all_files": [dict(f) for f in all_files],
        })
    db.close()

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def set_east_asia(el, font_name):
            rPr = el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr"); el.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), font_name)

        status_map = {"provided": "✓ 已提供", "partial": "△ 部分提供", "pending": "✗ 未提供", "na": "— 不适用"}
        risk_map = {"high": "高", "medium": "中", "low": "低"}

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        set_east_asia(normal._element, "Microsoft YaHei")

        # Title
        tp = doc.add_paragraph()
        tr = tp.add_run("创始人背景核查报告")
        tr.bold = True; tr.font.size = Pt(20)
        set_east_asia(tr._element, "Microsoft YaHei")

        pn = doc.add_paragraph()
        pr = pn.add_run(str(project["name"] or ""))
        pr.font.size = Pt(13)
        set_east_asia(pr._element, "Microsoft YaHei")

        from datetime import datetime
        doc.add_paragraph().add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10)
        doc.add_paragraph()

        for fd in founder_data:
            f = fd["founder"]
            # Founder heading
            fh = doc.add_paragraph()
            fhr = fh.add_run(f"▌ {f['name']}  {f.get('role') or ''}")
            fhr.bold = True; fhr.font.size = Pt(14)
            set_east_asia(fhr._element, "Microsoft YaHei")

            # group items by dimension
            from itertools import groupby
            dims_seen = []
            dims_map = {}
            for item in fd["items"]:
                d = item["dimension"]
                if d not in dims_map:
                    dims_map[d] = []
                    dims_seen.append(d)
                dims_map[d].append(item)

            for dim in dims_seen:
                dim_label = DIMENSION_LABELS.get(dim, {}).get("zh", dim)
                dp = doc.add_paragraph()
                dr = dp.add_run(f"  {dim}. {dim_label}")
                dr.bold = True; dr.font.size = Pt(11)
                set_east_asia(dr._element, "Microsoft YaHei")

                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for i, h in enumerate(["编号", "事项", "风险", "状态", "文件/说明"]):
                    hdr[i].text = h

                for item in dims_map[dim]:
                    row = tbl.add_row().cells
                    row[0].text = item["code"]
                    row[1].text = item["title"]
                    row[2].text = risk_map.get(item["risk_level"], "")
                    row[3].text = status_map.get(item["status"], item["status"])
                    parts = [f["file_name"] for f in item["files"]]
                    if item.get("statement"):
                        parts.append(item["statement"])
                    row[4].text = "\n".join(parts)

                doc.add_paragraph()

        # save docx to memory
        docx_buf = io.BytesIO()
        doc.save(docx_buf)
        docx_buf.seek(0)

        # build ZIP: docx + all founder files
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("创始人背景核查报告.docx", docx_buf.read())
            for fd in founder_data:
                fname = _sanitize(fd["founder"]["name"] or f"founder_{fd['founder']['id']}")
                for ff in fd["all_files"]:
                    src = ff.get("file_path", "")
                    if src and os.path.isfile(src):
                        zip_name = f"{fname}/{ff['file_name']}"
                        try:
                            zf.write(src, zip_name)
                        except Exception:
                            pass

        zip_buf.seek(0)
        proj_name = _sanitize(str(project["name"] or project_id))
        return StreamingResponse(
            zip_buf,
            media_type="application/zip",
            headers={"Content-Disposition": _content_disposition_attachment(f"founder_report_{proj_name}.zip")}
        )
    except Exception as e:
        raise HTTPException(500, f"报告生成失败：{str(e)}")
